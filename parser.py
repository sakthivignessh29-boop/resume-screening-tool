import os
import re

def extract_text(filepath):
    """Extract raw text from a resume file (PDF, DOCX, or TXT)."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.pdf':
        return _extract_pdf(filepath)
    elif ext in ('.docx', '.doc'):
        return _extract_docx(filepath)
    else:
        return _extract_txt(filepath)

def _extract_pdf(filepath):
    try:
        import pdfplumber
        text = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        return '\n'.join(text)
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {e}")

def _extract_docx(filepath):
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table cell text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return '\n'.join(paragraphs)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {e}")

def _extract_txt(filepath):
    try:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        raise ValueError(f"Failed to read TXT file: {e}")

def extract_candidate_name(text):
    """Attempt to extract the candidate's name from the first few lines."""
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    # The name is usually the first non-empty line that doesn't look like a header
    for line in lines[:5]:
        # Skip lines that look like contact info or headers
        if re.search(r'@|http|www\.|linkedin|github|resume|curriculum|vitae', line, re.IGNORECASE):
            continue
        # Skip lines with many special chars (section headers)
        if len(re.findall(r'[A-Za-z]', line)) < 3:
            continue
        # Name is typically 2-4 words, all title-case / caps
        words = line.split()
        if 2 <= len(words) <= 5 and all(w[0].isupper() for w in words if w.isalpha()):
            return line
    return lines[0] if lines else "Unknown"
